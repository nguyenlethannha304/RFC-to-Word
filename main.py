from cgitb import text
from gettext import textdomain
from itertools import chain
from operator import index
from PyPDF4.pdf import PdfFileReader
import os
import sys
import string
import re
from docx import Document
DIR_PATH = os.path.dirname(__file__)
NORMAL_LINE = r'^(\s{3}(\S{2}|(\S\s\S)))'
TABLE_LINE = r'^((\s{3})((\+-)|(\|\s)))'
HEADING_LINE = r'^(\d\.)+'
ORDER_LIST = r'^(\s{3}((\d\.)|(o\s)))'


class ConvertRfcToWord:
    def __init__(self, file_name, start_page, end_page):
        self.file_name = file_name
        self.file = None
        self.start_page = int(start_page)
        self.end_page = int(end_page) + 1
        self.state = ''
        self.pre_state = ''
        self.word_file = self.create_word_file()  # output file
        self.current_paragraph = self.word_file.add_paragraph()

    def check_page_valid(self):
        max_possible_page = self.file.getNumPages()
        # start_page must greater than 0, end_page must lesser or equal max_possible_page, and end_page >= start_page
        if self.start_page <= 0:
            raise "start_page: %s must be greater than 0" % (self.start_page)
        if self.end_page > max_possible_page:
            raise "end_page: %s must not larger than final page of document: %s" % (
                self.end_page - 1, max_possible_page)
        if self.start_page > self.end_page - 1:
            raise "start_page: %s must not greater than end_page: %s" % (
                self.start_page, self.end_page - 1)

    def open_file(self):
        self.file = PdfFileReader(open(f"{DIR_PATH}{self.file_name}", 'rb'))

    def create_word_file(self):
        return Document()

    def output_file(self):
        self.word_file.save(f"{DIR_PATH}output.docx")

    def __call__(self):
        self.open_file()
        self.check_page_valid()
        self.process()
        self.output_file()

    def process(self):
        for page_number in range(self.start_page, self.end_page):
            textList = self.get_text_list_from_page(page_number)
            textList = self.remove_header_footer_of_textList(textList)
            self.add_textList_to_word_file(textList)

    def get_text_list_from_page(self, page_number):
        page_object = self.file.getPage(page_number)
        text = page_object.extractText()
        textList = text.split('\n')
        return textList

    def remove_header_footer_of_textList(self, textList):
        # Header is first element of list and footer is last element of list
        textList.pop()
        del textList[0]
        return textList

    def add_textList_to_word_file(self, textList):
        for text_item in textList:
            if text_item == '':
                continue
            if self.is_heading(text_item):
                self.state = 'heading'
            elif self.is_order_list(text_item):
                self.state = 'order_list'
            elif self.is_table_line(text_item):
                self.state = 'table_line'
            elif self.is_normal_line(text_item):
                self.state = 'normal_line'
            else:
                self.state = 'other'
            text_item = text_item.strip()
            self.process_line(text_item)
            self.pre_state = self.state

    def is_heading(self, text_item):
        return re.match(HEADING_LINE, text_item)

    def is_order_list(self, text_item):
        return re.match(ORDER_LIST, text_item)

    def is_table_line(self, text_item):
        return re.match(TABLE_LINE, text_item)

    def is_normal_line(self, text_item):
        return re.match(NORMAL_LINE, text_item)

    def process_line(self, text_item):
        if self.state != self.pre_state:
            self.current_paragraph = self.word_file.add_paragraph()
        if self.state == 'heading':
            self.process_heading(text_item)
        elif self.state == 'table_line':
            self.process_table_line(text_item)
        elif self.state == 'normal_line':
            self.process_normal_line(text_item)
        elif self.state == 'order_list':
            self.process_order_list(text_item)
        else:
            self.process_other_case(text_item)

    def process_heading(self, text_item):
        if self.state == self.pre_state:
            self.current_paragraph = self.word_file.add_paragraph()
        index_end = 0
        while text_item[index_end] != ' ':
            index_end += 1
        heading_level = int(index_end / 2)
        self.current_paragraph.style = self.word_file.styles['Heading %s' % (
            heading_level)]
        self.current_paragraph.add_run(text_item)

    def process_order_list(self, text_item):
        if self.state == self.pre_state:
            self.current_paragraph = self.word_file.add_paragraph()
        self.current_paragraph.style = self.word_file.styles['List Paragraph']
        self.current_paragraph.add_run(text_item)

    def process_table_line(self, text_item):
        if self.current_paragraph.text != '':
            self.current_paragraph = self.word_file.add_paragraph()
        self.current_paragraph.style = self.word_file.styles['No Spacing']
        self.current_paragraph.add_run(text_item)

    def process_normal_line(self, text_item):
        if (self.current_paragraph.text != '' and self.current_paragraph.text[-1] in string.punctuation) and (text_item[0] not in string.ascii_lowercase):
            self.current_paragraph = self.word_file.add_paragraph()
        self.current_paragraph.style = self.word_file.styles['Normal']
        if self.pre_state == 'normal_line':
            text_item = ' ' + text_item
        self.current_paragraph.add_run(text_item)

    def process_other_case(self, text_item):
        if self.pre_state == 'order_list':
            self.current_paragraph.add_run(text_item)
            self.current_state = 'order_list'
        else:
            if self.current_paragraph.text == '':
                self.current_paragraph.add_run(text_item)
            else:
                self.current_paragraph = self.word_file.add_paragraph()
                self.current_paragraph.add_run(text_item)
        self.current_paragraph.style = self.word_file.styles['List Paragraph']


class Process:
    pass


if __name__ == "__main__":
    ConvertRfcToWord('rfc.pdf', 6, 47)()
